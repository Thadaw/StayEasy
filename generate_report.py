from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

IMG_DIR = r'E:\roshan\StayEasy\report-screenshots'

def add_image(name, width=Inches(5.5)):
    path = os.path.join(IMG_DIR, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        doc.add_paragraph(f'[Screenshot: {name} -- not available]')

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

# ── COVER PAGE ──
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('StayEasy')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Project Development Report')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('A multi-tenant Hotel and Restaurant Management SaaS platform\nbuilt with React, TypeScript, and Vite')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

for _ in range(4):
    doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('August 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ── TABLE OF CONTENTS (manual) ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Technology Stack',
    '2. Project Overview',
    '3. Host Landing Page',
    '4. Authentication System',
    '5. Property Onboarding Wizard',
    '    5.1 Property Type Selection',
    '    5.2 Property Details',
    '    5.3 Location Details & Maps',
    '    5.4 Photos & Amenities',
    '    5.5 Localization Settings',
    '    5.6 Branding & Visuals',
    '    5.7 Room Setup',
    '    5.8 Pricing & Offers',
    '    5.9 Review & Launch',
    '6. Dashboard & Analytics',
    '7. Management Pages',
    '8. Design System & Styling',
    '9. State Management',
    '10. API Integration Layer',
    '11. Internationalization',
    '12. Recent Improvements',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith('    '):
        p.paragraph_format.left_indent = Cm(1.5)

doc.add_page_break()

# ── 1. TECHNOLOGY STACK ──
doc.add_heading('1. Technology Stack', level=1)

doc.add_paragraph(
    'StayEasy is built on a modern frontend stack designed for performance, '
    'type safety, and developer experience. The application is a single-page '
    'app (SPA) that communicates with a remote backend hosted on Render.'
)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Technology'
hdr[1].text = 'Purpose'
techs = [
    ('React 18', 'UI library for building component-based interfaces'),
    ('TypeScript 5', 'Static type checking for JavaScript'),
    ('Vite 8', 'Fast build tool and dev server with HMR'),
    ('Tailwind CSS 4', 'Utility-first CSS framework for rapid styling'),
    ('Zustand 5', 'Lightweight global state management'),
    ('React Query 5', 'Server state management with caching and invalidation'),
    ('React Hook Form 7', 'Performant form handling with minimal re-renders'),
    ('Zod 4', 'Schema-based form validation'),
    ('Framer Motion', 'Animation library for page transitions and UI effects'),
    ('Leaflet + OpenStreetMap', 'Interactive maps with geocoding support'),
    ('Recharts', 'Composable charting library for dashboards'),
    ('React Router 7', 'Client-side routing with lazy-loaded routes'),
    ('Axios', 'HTTP client with interceptors for auth and error handling'),
    ('i18next', 'Internationalization framework (21 languages)'),
    ('Lucide React', 'Consistent icon set'),
    ('Stripe + Razorpay', 'Payment processing integrations'),
]
for tech, purpose in techs:
    row = table.add_row().cells
    row[0].text = tech
    row[1].text = purpose

doc.add_page_break()

# ── 2. PROJECT OVERVIEW ──
doc.add_heading('2. Project Overview', level=1)

doc.add_paragraph(
    'StayEasy is a hospitality property management system that lets property '
    'owners list, manage, and monetize their accommodations. The platform '
    'handles the full lifecycle from onboarding a new property through '
    'day-to-day operations including bookings, room management, pricing, '
    'and analytics.'
)

doc.add_paragraph(
    'The host-facing side of the application is the primary focus of this '
    'report. It includes a property onboarding wizard, dashboard analytics, '
    'and a suite of management pages for handling bookings, rooms, guests, '
    'staff, housekeeping, pricing, and reports.'
)

doc.add_heading('Key Capabilities', level=2)
features = [
    '9-step property onboarding wizard with auto-save and draft persistence',
    'Interactive maps with forward and reverse geocoding',
    'Real-time dashboard with revenue, occupancy, and booking analytics',
    'Multi-step room configuration with per-room amenities and policies',
    'Dynamic pricing with seasonal rates, discounts, and custom offers',
    'Dual-role authentication (host and guest) with OTP verification',
    '21-language internationalization with RTL support',
    'Responsive design with dark mode support',
    'Payment integration via Stripe and Razorpay',
]
for f in features:
    add_bullet(f)

doc.add_page_break()

# ── 3. HOST LANDING PAGE ──
doc.add_heading('3. Host Landing Page', level=1)

doc.add_paragraph(
    'The host landing page serves as the marketing entry point for property '
    'owners. It features a hero section with the headline "Your place, your '
    'terms, zero cut" alongside animated card elements showing payout receipts, '
    'guest messages, and listing cards. A scrolling ticker strip highlights '
    'key benefits like no listing fees, set-your-own pricing, and weekly payouts.'
)

doc.add_paragraph(
    'Below the fold, there are three numbered feature cards explaining the '
    'platform benefits, a gold-themed trust section highlighting verified '
    'properties and 24/7 support, a three-step onboarding flow (Open account '
    'in ~2 min, Publish listing in ~10 min, Get paid weekly), and an FAQ '
    'accordion with six expandable questions. The page uses Framer Motion '
    'for entrance animations and scroll-triggered effects.'
)

add_image('01-landing-page.png')

doc.add_page_break()

# ── 4. AUTHENTICATION SYSTEM ──
doc.add_heading('4. Authentication System', level=1)

doc.add_paragraph(
    'The authentication system supports two roles: host and guest. Each role '
    'has its own registration flow, session management, and protected routes. '
    'The login page uses a split-panel design with an animated video on the '
    'left and the login form on the right.'
)

doc.add_heading('Login Flow', level=2)
doc.add_paragraph(
    'The login form collects email and password, with a "Remember me" toggle '
    'that determines whether the session persists in localStorage or '
    'sessionStorage. The form uses OAuth2 password grant (form-urlencoded) '
    'to authenticate against the backend. On success, a JWT token is stored '
    'and an Axios interceptor attaches it to all subsequent API requests.'
)

add_image('02-login-page.png')

doc.add_heading('Registration Flow', level=2)
doc.add_paragraph(
    'The signup form collects full name, phone number, email, and password. '
    'For guests, nationality is also collected. After submission, a 6-digit '
    'OTP is sent to the email address. The user enters the code, and upon '
    'verification, the account is activated. The OTP input uses individual '
    'digit fields with auto-advance on entry. A 30-second cooldown timer '
    'prevents spamming the resend button.'
)

doc.add_heading('Password Requirements', level=2)
doc.add_paragraph(
    'Passwords must be at least 8 characters long and include at least one '
    'number and one special character. The signup form shows real-time '
    'validation feedback as the user types.'
)

doc.add_heading('Session Management', level=2)
doc.add_paragraph(
    'The AuthContext provider manages token persistence, expiry tracking '
    '(30-day window), and user normalization. The backend sometimes returns '
    'snake_case and sometimes camelCase fields, so the auth layer normalizes '
    'these inconsistencies. On logout, all draft data is cleaned up from '
    'localStorage. A 401 response from any API call triggers an automatic '
    'redirect to the appropriate login page with a return URL.'
)

doc.add_page_break()

# ── 5. PROPERTY ONBOARDING WIZARD ──
doc.add_heading('5. Property Onboarding Wizard', level=1)

doc.add_paragraph(
    'The property onboarding wizard is the centerpiece of the host portal. '
    'It guides property owners through a 9-step process to list their '
    'property on the platform. Each step saves data to the backend '
    'immediately, and the entire wizard state is persisted to localStorage '
    'every 500ms so users can resume where they left off even if they '
    'close the browser.'
)

doc.add_paragraph(
    'The wizard is organized into three high-level sections shown on the '
    'progress bar: Property Details (steps 1-5), Room Details (step 6), '
    'and Pricing & Offers (steps 7-9). The progress bar is clickable, '
    'allowing users to jump back to any completed section directly.'
)

# 5.1
doc.add_heading('5.1 Property Type Selection', level=2)
doc.add_paragraph(
    'The first step presents eight property types in a grid: Hotel, Resort, '
    'Restaurant, Hostel, Villa, Apartment, Guesthouse, and Other. Each type '
    'has an icon from the Lucide icon set. Selecting a type automatically '
    'advances to the next step. The selected type influences subsequent '
    'wizard options and how the property appears in search results.'
)

add_image('03-step1-property-type.png')

# 5.2
doc.add_heading('5.2 Property Details', level=2)
doc.add_paragraph(
    'The property details form collects the fundamental information about '
    'the property: name, total number of rooms, number of floors (with '
    'increment/decrement controls), year built, a detailed description '
    '(up to 2500 characters with a live counter), phone number, and '
    'official email. The description field includes helper text suggesting '
    'what to highlight: best features, neighborhood vibes, and recent '
    'renovations.'
)

add_image('04-step2-property-details.png')

# 5.3
doc.add_heading('5.3 Location Details & Maps', level=2)
doc.add_paragraph(
    'The location step features an embedded Leaflet map with OpenStreetMap '
    'tiles. Users can either type an address (country, state, city, ZIP, '
    'street) and the map updates to show that location, or click directly '
    'on the map and the address fields are auto-filled via reverse geocoding '
    'through the Nominatim API. This bidirectional syncing means the map '
    'and form are always in sync. A custom SVG marker pinpoints the exact '
    'location on the map.'
)

doc.add_paragraph(
    'The country dropdown includes 12 options: United States, United Kingdom, '
    'Canada, Australia, India, Nepal, Germany, France, Italy, Spain, Japan, '
    'and Brazil. The map supports zoom and pan for precise positioning.'
)

add_image('05-step3-location.png')

# 5.4
doc.add_heading('5.4 Photos & Amenities', level=2)
doc.add_paragraph(
    'The photos step provides a drag-and-drop upload zone that accepts up '
    'to five images. Uploaded photos appear in a preview grid where users '
    'can reorder them and select a cover photo by clicking the star button '
    'on any image. Each photo has a remove button.'
)

doc.add_paragraph(
    'Below the photo section, users set a star rating (1-5) using clickable '
    'star icons. The amenities section fetches available amenities from the '
    'backend API and displays them with a search filter. Users select system '
    'amenities by clicking on them, and can also add custom amenities with '
    'a text input and add button.'
)

# 5.5
doc.add_heading('5.5 Localization Settings', level=2)
doc.add_paragraph(
    'The localization step configures region-specific settings: currency '
    '(9 options including USD, EUR, INR, NPR, GBP, JPY, AUD, CAD, BRL), '
    'timezone (9 options covering major global zones), and language preference '
    '(English US, English UK, or Spanish as selectable pills).'
)

doc.add_paragraph(
    'Check-in and check-out times are configurable with text inputs, and '
    'early/late grace period dropdowns allow setting flexible arrival and '
    'departure windows. A toggle for "Always Allow Check-in" bypasses '
    'time restrictions when enabled.'
)

# 5.6
doc.add_heading('5.6 Branding & Visuals', level=2)
doc.add_paragraph(
    'The branding step lets property owners customize their listing appearance. '
    'A logo upload shows a preview card with the property name and phone '
    'number rendered with the selected branding. The brand color picker '
    'includes a native color input, a hex code input field, and five '
    'suggested palette swatches for quick selection.'
)

doc.add_paragraph(
    'A WCAG 2.1 compliance badge automatically checks whether the selected '
    'brand color passes accessibility contrast requirements against white '
    'text. A live phone-frame preview shows how the property listing will '
    'appear on a mobile device with the chosen branding applied.'
)

# 5.7
doc.add_heading('5.7 Room Setup', level=2)
doc.add_paragraph(
    'The room setup step is the most detailed part of the wizard. Each room '
    'is configured independently with: floor number, room name, room type '
    '(5 options), bed type (5 options), photo upload (up to 5 per room with '
    'cover selection), occupancy settings (max adults, max children, pets '
    'allowed toggle), base rate per night, and cancellation policy.'
)

doc.add_paragraph(
    'The cancellation policy offers four presets (Flexible, Moderate, Strict, '
    'Non-refundable) plus a custom option where users write their own title '
    'and description. Custom policies can be saved and reused across rooms. '
    'Each room also has its own amenity selection pulled from the system '
    'amenities list with a search filter.'
)

doc.add_paragraph(
    'Two shortcut buttons speed up the process: "Copy Last Room" duplicates '
    'the most recently configured room, and "Add Blank Room" creates a new '
    'empty room card. Rooms can be expanded or collapsed for easier navigation.'
)

# 5.8
doc.add_heading('5.8 Pricing & Offers', level=2)
doc.add_paragraph(
    'The pricing step manages special offers and promotional discounts. Four '
    'preset offers are available out of the box: Early Bird Discount (10% off '
    'for bookings 30+ days in advance), Last-Minute Deal (15% off for bookings '
    'within 48 hours), Long Stay Discount (20% off for 7+ night stays), and '
    'Free Cancellation (full refund if cancelled 48+ hours before check-in).'
)

doc.add_paragraph(
    'Each offer has a toggle switch to enable/disable it, and a date range '
    'picker (using react-datepicker) to set the active period. A "Create '
    'Custom Offer" button opens a modal where users can define their own '
    'promotional offer with a title, description, discount percentage, and '
    'date range.'
)

# 5.9
doc.add_heading('5.9 Review & Launch', level=2)
doc.add_paragraph(
    'The final step provides a comprehensive summary of all entered data. '
    'Each section (Core Identity, Media, Address, Compliance) has an edit '
    'button that jumps back to the relevant wizard step. A profile strength '
    'bar calculates a completion percentage based on how many fields have '
    'been filled. Room cards display all configured room details. The '
    '"Publish Listing" button activates the property on the platform by '
    'calling the activation endpoint, then clears the draft from localStorage.'
)

doc.add_page_break()

# ── 6. DASHBOARD & ANALYTICS ──
doc.add_heading('6. Dashboard & Analytics', level=1)

doc.add_paragraph(
    'The overall dashboard provides a portfolio-wide view of all properties. '
    'It displays four key metric cards: Total Revenue, Occupancy Rate, Total '
    'Bookings, and Active Properties. Each card shows the current value with '
    'a percentage change indicator comparing to the previous period.'
)

doc.add_paragraph(
    'Below the metrics, a revenue bar chart (built with Recharts) compares '
    'this month, last month, and this week side by side. An occupancy pie '
    'chart shows the breakdown of room statuses: occupied, available, '
    'maintenance, and out-of-order. Additional sections show recent bookings, '
    'room status summary, and quick action shortcuts.'
)

add_image('12-dashboard.png')

doc.add_paragraph(
    'The per-property dashboard mirrors the overall layout but scopes all '
    'metrics to a single property. It includes additional sections for '
    'arrivals/departures and a restaurant overview module.'
)

doc.add_page_break()

# ── 7. MANAGEMENT PAGES ──
doc.add_heading('7. Management Pages', level=1)

doc.add_paragraph(
    'Beyond the wizard and dashboard, the host portal includes a full suite '
    'of operational management pages accessible from the sidebar navigation.'
)

doc.add_heading('Sidebar Navigation', level=2)
doc.add_paragraph(
    'The sidebar is organized into two sections: MAIN (Dashboard, Property '
    'Management, Bookings, Guests, Staff, Housekeeping, Reports) and SYSTEM '
    '(Settings, Notifications, Activity, Support). The sidebar is collapsible, '
    'has badge counts on relevant items, and includes a search filter for '
    'quick navigation.'
)

doc.add_heading('Bookings Management', level=2)
doc.add_paragraph(
    'The bookings page displays booking statistics, filters by status, room '
    'type, and date range, and shows a table of all bookings with details '
    'like guest name, room type, check-in/out dates, amount, and status.'
)

doc.add_heading('Room Management', level=2)
doc.add_paragraph(
    'The rooms page provides a grid or table view of all rooms with their '
    'current status, type, floor, rate, and occupancy. Rooms can be filtered '
    'by status, floor, and type.'
)

doc.add_heading('Staff & Housekeeping', level=2)
doc.add_paragraph(
    'The staff page manages employee information with a table view, filters, '
    'and pagination. The housekeeping page handles task management with '
    'assignable tasks, room status tracking, and staff assignment views.'
)

doc.add_heading('Pricing Management', level=2)
doc.add_paragraph(
    'The pricing page has four sub-views: Overview (stats, feature cards, '
    'recent activity, upcoming promotions), Seasonal Pricing (timeline view '
    'and data table), Discounts & Offers (offer list with detail panel), '
    'and Packages (package management with detail view).'
)

doc.add_heading('Reports', level=2)
doc.add_paragraph(
    'The reports page includes revenue summaries, charts, revenue breakdown '
    'by department, top room types analysis, and recent bookings data. '
    'Filters and date range selection allow drilling into specific periods.'
)

doc.add_heading('Settings', level=2)
doc.add_paragraph(
    'The settings page is organized into tabs: Branding, Users & Roles, '
    'Payment Gateway, Localization, Security, and Audit Logs. Each tab '
    'provides configuration options for the respective area.'
)

doc.add_page_break()

# ── 8. DESIGN SYSTEM ──
doc.add_heading('8. Design System & Styling', level=1)

doc.add_paragraph(
    'The application uses a comprehensive CSS custom properties design '
    'system defined in theme.css. This provides a single source of truth '
    'for colors, spacing, typography, and other visual tokens.'
)

doc.add_heading('Color Palette', level=2)
colors = [
    ('Primary', '#1A3C5E', 'Deep navy blue used for headings and key UI elements'),
    ('Accent', '#2E86AB', 'Teal blue for interactive elements and highlights'),
    ('Success', '#1E8449', 'Green for positive states and confirmations'),
    ('Warning', '#D35400', 'Orange for caution states'),
    ('Danger', '#C0392B', 'Red for errors and destructive actions'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Token'
hdr[1].text = 'Hex Value'
hdr[2].text = 'Usage'
for name, hex_val, usage in colors:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = hex_val
    row[2].text = usage

doc.add_heading('Typography', level=2)
doc.add_paragraph(
    'The design system uses four font families: Sora for headings, Inter '
    'for body text, Playfair Display for serif accents, and Plus Jakarta '
    'Sans as an alternative. Font sizes follow a modular scale from 12px '
    'to 36px.'
)

doc.add_heading('Dark Mode', level=2)
doc.add_paragraph(
    'Full dark mode support is implemented through a .dark class variant '
    'on the root element. All semantic tokens (background, foreground, card, '
    'muted, accent, border, input, ring) have dark mode equivalents.'
)

doc.add_heading('Responsive Breakpoints', level=2)
breakpoints = [
    ('Mobile', '375px'),
    ('Tablet', '768px'),
    ('Desktop', '1280px'),
    ('Wide', '1440px'),
]
for name, bp in breakpoints:
    add_bullet(f'{name}: {bp}')

doc.add_page_break()

# ── 9. STATE MANAGEMENT ──
doc.add_heading('9. State Management', level=1)

doc.add_paragraph(
    'The application uses a combination of state management approaches '
    'depending on the type of state:'
)

doc.add_heading('Zustand Stores', level=2)
doc.add_paragraph(
    'Two lightweight Zustand stores handle global UI state: uiStore manages '
    'sidebar collapsed state, and propertyStore tracks the currently selected '
    'property ID (persisted to localStorage).'
)

doc.add_heading('React Context', level=2)
doc.add_paragraph(
    'Four React Context providers manage application-wide state: AuthProvider '
    '(authentication state, login/signup/logout), FavoritesProvider (favorite '
    'properties with localStorage persistence), BookingProvider (booking '
    'state), and CouponProvider (coupon management).'
)

doc.add_heading('React Query', level=2)
doc.add_paragraph(
    'React Query (TanStack Query) handles all server state with organized '
    'query key factories for properties, rooms, room types, bed types, '
    'special offers, discount codes, tenants, and bookings. The overall '
    'dashboard uses useQueries for parallel data fetching across multiple '
    'endpoints. Mutations use invalidateQueries for automatic cache '
    'invalidation.'
)

doc.add_page_break()

# ── 10. API INTEGRATION ──
doc.add_heading('10. API Integration Layer', level=1)

doc.add_paragraph(
    'The API service layer (pmsApi.ts) wraps 30+ backend endpoints organized '
    'into categories: Properties, Location, Photos/Amenities, Localization, '
    'Branding, Images, Rooms, Room Types, Bed Types, Special Offers, Discount '
    'Codes, Tenants, and Bookings.'
)

doc.add_paragraph(
    'An unwrapBody helper function handles both envelope-wrapped responses '
    '(where the actual data is inside a "data" field) and bare JSON responses, '
    'so the frontend works seamlessly with either backend format.'
)

doc.add_paragraph(
    'The Axios instance (axios.ts) auto-attaches Bearer tokens to all requests, '
    'handles 401 responses by clearing the session and redirecting to login, '
    'and provides a skip-auth-redirect option for the initial /me handshake '
    'call.'
)

doc.add_heading('Backend Proxy', level=2)
doc.add_paragraph(
    'An Express.js server (server.mjs) runs on port 3001 and serves two '
    'purposes: in development, it proxies /api requests to the remote Render '
    'backend to avoid CORS issues; in production, it handles Razorpay payment '
    'order creation/verification and Stripe webhook events for '
    'payment_intent.succeeded and payment_intent.payment_failed.'
)

doc.add_page_break()

# ── 11. I18N ──
doc.add_heading('11. Internationalization', level=1)

doc.add_paragraph(
    'The application supports 21 languages through i18next and '
    'react-i18next. The supported languages are: English, Nepali, Hindi, '
    'Spanish, French, German, Italian, Portuguese, Dutch, Japanese, Korean, '
    'Chinese, Arabic (with RTL support), Russian, Thai, Vietnamese, Turkish, '
    'Polish, Swedish, and Danish.'
)

doc.add_paragraph(
    'When Arabic is selected, the application automatically toggles RTL '
    '(right-to-left) direction on the document. Locale files are stored '
    'in src/locales/ with separate JSON files for each language.'
)

doc.add_page_break()

# ── 12. RECENT IMPROVEMENTS ──
doc.add_heading('12. Recent Improvements', level=1)

doc.add_paragraph(
    'The following improvements were implemented during the current development '
    'cycle:'
)

doc.add_heading('Clickable Progress Bar', level=2)
doc.add_paragraph(
    'The progress bar in the property onboarding wizard was upgraded from a '
    'static display to an interactive navigation element. Users can now click '
    'on any completed section to jump directly to that step. Completed steps '
    'show a green checkmark and respond to hover with a scale animation and '
    'dotted underline on the label. Future steps remain non-clickable.'
)

doc.add_heading('Smart Re-save Prevention', level=2)
doc.add_paragraph(
    'A savedSteps tracking system was added to prevent duplicate backend records '
    'when navigating back and forward through the wizard. The system tracks '
    'which steps have been saved and takes data snapshots after each save. '
    'When a user goes back, makes no changes, and clicks Next again, the API '
    'call is skipped entirely. If changes are detected (via JSON comparison), '
    'only then is the save triggered.'
)

doc.add_heading('Draft Persistence Enhancements', level=2)
doc.add_paragraph(
    'The localStorage-based draft system was refined with user-scoped keys '
    '(serveIQDraft_{userId}), 500ms debounced auto-save, and proper cleanup '
    'on logout and successful publish. The draft state includes the current '
    'step, all form data, and the property ID for resuming backend operations.'
)

doc.add_page_break()

# ── CONCLUSION ──
doc.add_heading('Conclusion', level=1)

doc.add_paragraph(
    'StayEasy is a comprehensive hospitality management platform with a '
    'polished host portal that handles the complete property lifecycle from '
    'onboarding through daily operations. The 9-step property wizard with '
    'draft persistence, interactive maps, real-time branding preview, and '
    'smart re-save logic represents the most complex piece of the '
    'application. The dashboard and management suite provide the tools needed '
    'to run day-to-day operations effectively.'
)

doc.add_paragraph(
    'The codebase is organized with clear separation of concerns: feature '
    'slices for property and booking logic, a centralized API layer, '
    'component-based UI architecture, and a design system that ensures '
    'visual consistency across all pages. The use of TypeScript throughout '
    'provides type safety and better developer experience.'
)

# Save
output_path = r'E:\roshan\StayEasy\StayEasy_Project_Report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
